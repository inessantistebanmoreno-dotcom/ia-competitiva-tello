# =============================================================================
# crear_memoria.py — Genera la memoria del proyecto en Word (.docx)
# Ejecutar con: python crear_memoria.py
# =============================================================================

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

TELLO_RED = RGBColor(0xBA, 0x0C, 0x2F)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x47, 0x55, 0x69)
GREEN = RGBColor(0x16, 0xA3, 0x4A)

doc = Document()

# ── Márgenes ──
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

# ── Estilos base ──
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)
style.font.color.rgb = DARK

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = TELLO_RED
    run.font.name = 'Arial'
    # Línea bajo el título
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), 'BA0C2F')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = DARK
    run.font.name = 'Arial'
    return p

def heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = TELLO_RED
    run.font.name = 'Arial'
    return p

def body(text, bold_parts=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.5)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.name = 'Arial'
        r1.font.size = Pt(11)
        r2 = p.add_run(text)
        r2.font.name = 'Arial'
        r2.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(11)
    return p

def add_table(headers, rows, col_widths_cm=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Background color
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'BA0C2F')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)
    # Data rows
    for idx, row_data in enumerate(rows):
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row_cells[i].text = str(cell_text)
            run = row_cells[i].paragraphs[0].runs[0]
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            # Alternate shading
            if idx % 2 == 1:
                tc = row_cells[i]._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), 'F8FAFC')
                shd.set(qn('w:val'), 'clear')
                tcPr.append(shd)
    # Column widths
    if col_widths_cm:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                if i < len(col_widths_cm):
                    cell.width = Cm(col_widths_cm[i])
    doc.add_paragraph()
    return table

# ═══════════════════════════════════════════════════════════
# PORTADA
# ═══════════════════════════════════════════════════════════
p_title = doc.add_paragraph()
p_title.paragraph_format.space_before = Pt(60)
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_title.add_run('MEMORIA DEL PROYECTO')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = TELLO_RED
run.font.name = 'Arial'

p_subtitle = doc.add_paragraph()
p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_subtitle.add_run('Agente de Inteligencia Competitiva de Etiquetado')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = DARK
run.font.name = 'Arial'

doc.add_paragraph()

p_company = doc.add_paragraph()
p_company.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_company.add_run('TELLO · Departamento de Marketing · Etiquetado')
run.font.size = Pt(12)
run.font.color.rgb = GRAY
run.font.name = 'Arial'

p_author = doc.add_paragraph()
p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_author.add_run('Inés Santisteban Moreno')
run.font.size = Pt(12)
run.font.color.rgb = GRAY
run.font.name = 'Arial'

p_date = doc.add_paragraph()
p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_date.add_run('Junio – Julio 2026')
run.font.size = Pt(11)
run.font.color.rgb = GRAY
run.font.name = 'Arial'
run.italic = True

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 1. DESCRIPCIÓN DEL PROYECTO
# ═══════════════════════════════════════════════════════════
heading1('1. DESCRIPCIÓN DEL PROYECTO')

body('El Agente de Inteligencia Competitiva de Etiquetado es un sistema automatizado desarrollado para Tello que monitoriza de forma continua el etiquetado nutricional, los ingredientes y los claims de packaging de los principales competidores del sector cárnico español: Noel, Campofrío, El Pozo y Argal.')
body('El sistema extrae automáticamente información de las webs de los competidores, la almacena en una base de datos estructurada y alerta al equipo de etiquetado cuando detecta cualquier modificación relevante, clasificando los cambios por su impacto en la calidad del producto.')

heading2('1.1 Objetivos')
bullet('Automatizar el seguimiento de 881 productos de 4 competidores')
bullet('Detectar cambios en ingredientes, tabla nutricional y claims en tiempo casi real')
bullet('Clasificar los cambios según su impacto: MEJORA o EMPEORAMIENTO de calidad')
bullet('Enviar alertas automáticas por email al equipo de etiquetado')
bullet('Proporcionar una interfaz visual para consultar y comparar productos')

heading2('1.2 Tecnologías utilizadas')
bullet('Python 3.14 + Playwright: scraping web y automatización del navegador')
bullet('playwright-stealth: acceso a webs con protección anti-bot (Argal)')
bullet('PostgreSQL / Supabase: base de datos en la nube (gratuita)')
bullet('GitHub Actions: automatización en la nube sin coste (repositorio público)')
bullet('SendGrid: envío de emails desde la nube (100 emails/día gratuitos)')
bullet('HTML / JavaScript + Supabase SDK: dashboard visual interactivo')

heading2('1.3 Competidores monitorizados')
add_table(
    ['Competidor', 'Web', 'Productos', 'Técnica de extracción'],
    [
        ['Noel',       'noel.es',       '135', 'HTML directo, categorías /categorias/'],
        ['Campofrío',  'campofrio.es',  '105', 'SPA JavaScript, URLs /p/nombre'],
        ['El Pozo',    'elpozo.com',    '493', 'Paginación 42 páginas, /productos/page/N/'],
        ['Argal',      'argal.com',     '148', 'playwright-stealth (anti-reCAPTCHA)'],
        ['TOTAL',      '',              '881', ''],
    ],
    col_widths_cm=[3, 3.5, 2.5, 7]
)

# ═══════════════════════════════════════════════════════════
# 2. DIARIO DE TRABAJO
# ═══════════════════════════════════════════════════════════
doc.add_page_break()
heading1('2. DIARIO DE TRABAJO')

# ── Semana 1 ──
heading2('2.1 Semana 1 — Análisis y Diseño (26–31 Mayo 2026)')

heading3('26 de mayo de 2026 — Reunión de etiquetado')
body('Reunión inicial con el equipo de etiquetado donde se presenta la necesidad del proyecto. Se analizan los documentos existentes: Excel de productos Noel (estructura de referencia con 31 columnas), presentación del proyecto y documento de sistema de inteligencia competitiva.')
body('Se identifican los competidores a monitorizar y se definen los datos prioritarios a extraer:')
bullet('CRÍTICO: Ingredientes completos, tabla nutricional (kcal, proteínas, grasas, sal), alérgenos')
bullet('ALTO: Claims del packaging (% carne, sellos, certificaciones, gama)')
bullet('MEDIO: Categoría, subcategoría, formato, gramaje')
bullet('INFORMATIVO: URL del producto, descripción web, foto')

heading3('27–31 de mayo — Diseño técnico')
body('Se diseña la arquitectura del sistema en 4 bloques: Extracción (Scraping + IA), Base de datos (PostgreSQL/Supabase), Orquestación y Dashboard visual.')
body('Se diseña el esquema de la base de datos con dos tablas principales:')
bullet('productos_competencia: estado actual de cada producto (1 fila por referencia)')
bullet('historial_cambios: registro inmutable de cada cambio detectado con valor anterior y nuevo')
body('Se define el sistema de alertas por severidad:')
bullet('CRÍTICA: cambios en ingredientes o alérgenos → email inmediato al equipo')
bullet('ALTA: cambios en valores nutricionales → email inmediato')
bullet('MEDIA: nuevos productos detectados → resumen diario')
bullet('BAJA: cambios en claims o descripción → informe semanal')

# ── Semana 2 ──
doc.add_page_break()
heading2('2.2 Semana 2 — Desarrollo e Implementación (2–6 Junio 2026)')

heading3('2 de junio — Configuración del entorno')
body('Se instala Python 3.14 y se configura VS Code. Se resuelven problemas de compatibilidad con la versión de Python y las librerías (greenlet, playwright). Se instala Playwright con Chromium (~150 MB).')
body('Se crea la base de datos en Supabase (plan gratuito, servidor en Frankfurt, Alemania). Se ejecuta el schema SQL creando las tablas y las vistas necesarias para el dashboard. Se configura el archivo .env con las credenciales.')

heading3('2 de junio — Scraper de Noel')
body('Se desarrolla el scraper de Noel (noel.es). La web carga contenido con JavaScript, por lo que fue necesario identificar las URLs reales de categorías (/categorias/cocidos/, /categorias/curados/, /categorias/carne-fresca/, etc.).')
body('Resultado: 135 productos de Noel extraídos y guardados en la base de datos en el primer ciclo completo.')

heading3('3 de junio — Scraper de Campofrío')
body('Se desarrolla el scraper de Campofrío (campofrio.es). La web es una SPA (Single Page Application) con URLs de producto en formato /p/nombre-producto. Se identifican 12 categorías de productos.')
body('Resultado: 105 productos de Campofrío extraídos correctamente.')

heading3('3 de junio — Scraper de El Pozo')
body('Se desarrolla el scraper de El Pozo (elpozo.com). La web tiene paginación con 42 páginas y 496 productos. Se detecta que la tabla nutricional está en formato imagen (no texto), lo que impide la extracción automática sin la API de Claude Vision.')
body('Se reduce el delay entre páginas de 2 segundos a 0,8 segundos para optimizar el tiempo de ejecución en la nube.')
body('Resultado: 493 productos extraídos.')

heading3('3 de junio — Scraper de Argal')
body('Se intenta el scraper de Argal (argal.com). La web usa reCAPTCHA que bloquea los navegadores automatizados. Se resuelve instalando playwright-stealth, una librería que hace que el navegador automatizado parezca un usuario humano real.')
body('Se identifican las 11 categorías de productos: jamón cocido, pavo y pollo, mortadela, fuet, jamón y lomo curado, chorizo, ibéricos, salchichas, patés, snacks y platos preparados.')
body('Resultado: 148 productos de Argal extraídos correctamente. El reCAPTCHA ya no representa un obstáculo.')

heading3('3–4 de junio — Sistema de detección de cambios')
body('Se implementa el sistema de detección de cambios en database/db.py:')
bullet('Hash SHA-256 del contenido de cada producto para detectar cambios rápidamente')
bullet('Comparación campo por campo para identificar exactamente qué cambió')
bullet('Registro en historial_cambios con valor anterior y nuevo')
body('Se implementa el motor de calidad (calidad.py) basado en normativa europea (Reg. 1924/2006 y RD 474/2014). Cada cambio detectado se clasifica automáticamente como:')
bullet('MEJORA: bajada de sal, subida de % carne, eliminación de nitritos, etc.')
bullet('EMPEORAMIENTO: bajada de % carne, subida de sal, adición de carragenina, etc.')
bullet('NEUTRO: cambios de descripción web, actualización de imágenes')

heading3('4–5 de junio — Sistema de alertas por email')
body('Se implementa alertas_scheduler.py con dos tipos de envío automático:')
bullet('Alerta de nuevos productos: consulta el historial cada hora y envía email si hay novedades')
bullet('Informe diario: resumen de todos los cambios de las últimas 24 horas, ordenado por impacto en calidad (primero empeoramientos críticos, luego mejoras)')
body('Se configura el envío SMTP con Gmail y contraseña de aplicación. Se configuran las tareas en el Programador de Tareas de Windows con ejecución diaria a las 9:15.')

heading3('5–6 de junio — Dashboard visual')
body('Se desarrolla el dashboard HTML conectado a Supabase mediante JavaScript. Pantallas implementadas:')
bullet('Dashboard: KPIs, alertas recientes, gráfico de actividad por mes')
bullet('Alertas: lista filtrable por severidad (Crítica, Alta, Nueva, Baja)')
bullet('Productos: catálogo con panel lateral de filtros (marca, categoría, tipo carne, dieta)')
bullet('Comparativa: tabla nutricional comparada entre marcas, ordenada por % carne descendente')
bullet('Historial de cambios: registro cronológico con preview expandible de cada cambio')
bullet('Mapa de claims: matriz de presencia de claims por marca y categoría')
bullet('Análisis nutricional: comparativa por categoría (cocidos, curados, embutidos, pavo/pollo...)')
bullet('Vistas por competidor: Campofrío, El Pozo, Noel, Argal')

body('Total de productos en base de datos al final de la semana 2: 881 productos (135 Noel + 105 Campofrío + 493 El Pozo + 148 Argal).')

# ── Semana 3 ──
doc.add_page_break()
heading2('2.3 Semana 3 — Automatización en la Nube (7–9 Junio 2026)')

heading3('8 de junio — GitHub Actions')
body('Se crea repositorio público en GitHub: github.com/inessantistebanmoreno-dotcom/ia-competitiva-tello')
body('Se configura GitHub Actions para ejecutar el scraping automáticamente en los servidores de GitHub, sin necesidad de tener el ordenador encendido. Horario configurado:')
bullet('Ciclo de scraping: 8:00 y 16:00 (hora española)')
bullet('Informe diario por email: 9:15 (hora española)')
bullet('Coste: gratuito (repositorio público, minutos ilimitados)')
body('Primer test exitoso: scraping completo en 53 minutos y 47 segundos (Noel 14 min, Campofrío 14 min, El Pozo 1 min con delay reducido, Argal 22 min).')

heading3('8–9 de junio — SendGrid para emails desde la nube')
body('Gmail bloquea los envíos desde servidores en la nube por seguridad. Se configura SendGrid (servicio de email con 100 emails/día gratuitos) para el envío desde GitHub Actions.')
body('El email automático con el informe diario llega correctamente a marketing1@tello.es desde la nube, sin necesidad de tener el ordenador encendido.')

heading3('9 de junio — Mejoras al dashboard')
body('Se rediseña la pantalla de Productos con un panel lateral de filtros tipo acordeón:')
bullet('Filtros por Marca con checkboxes y contador de productos')
bullet('Filtros por Categoría con opción "Ver más" para categorías secundarias')
bullet('Filtros por Tipo de carne (cerdo, pavo, pollo, vacuno, ibérico...)')
bullet('Filtros por Dieta (Sin Gluten, Sin Lactosa, Sin Colorantes, Sin Conservantes)')
bullet('Chips activos que se pueden eliminar individualmente')
bullet('Ordenación configurable: Relevancia, % Carne, Sal, Nombre A-Z, Más reciente')
body('Se implementa diseño responsive para que el dashboard se vea correctamente en cualquier tamaño de pantalla sin necesidad de reducir el zoom.')

heading3('9 de junio — Análisis experto de ingredientes')
body('Se implementa lógica experta basada en normativa de etiquetado alimentario:')
body('1. Extracción correcta del % de carne: identifica todos los ingredientes cárnicos en la lista de ingredientes, extrae sus porcentajes individuales y los suma. Casos contemplados:')
bullet('"Jamón de cerdo (93%), agua, sal..." → 93%')
bullet('"Carne de cerdo (50%), carne de vacuno (35%), agua..." → 85% (suma)')
bullet('"Jamón de cerdo 50% raza Duroc (93%), agua..." → 93% (ignora el 50% intermedio)')
bullet('"Pechuga de pavo (60%), agua (20%), almidón (5%)..." → 60% (ignora no-cárnicos)')
body('2. Análisis funcional de ingredientes: clasifica automáticamente cada ingrediente por su función real, independientemente del nombre técnico usado por el fabricante:')
bullet('Agua añadida: agua, caldo, salmuera, agua de cocción')
bullet('Almidones/Rellenos: fécula de patata, almidón, maltodextrina, harina')
bullet('Azúcares: dextrosa, glucosa, jarabe de glucosa, sacarosa, lactosa')
bullet('Sal/Sodio: sal, cloruro sódico, nitrito sódico, nitrato potásico')
bullet('Derivados lácteos: lactosa, leche en polvo, proteínas de la leche, caseinato')
bullet('Conservadores: nitrito sódico, lactato potásico, eritorbato sódico')
bullet('Estabilizantes/Ligantes: tripolifosfato, carragenanos, goma guar, sorbitol')
bullet('Aromas: aroma, aroma natural, aroma de humo')
body('3. Corrección de falsos positivos numéricos: 2,10 y 2,1 ya no se detectan como cambio.')

# ═══════════════════════════════════════════════════════════
# 3. ESTADO ACTUAL
# ═══════════════════════════════════════════════════════════
doc.add_page_break()
heading1('3. ESTADO ACTUAL DEL PROYECTO (9 Junio 2026)')

heading2('3.1 Funcionalidades operativas')
add_table(
    ['Módulo', 'Estado', 'Detalle'],
    [
        ['Scraping Noel',           'Operativo', '135 productos, extracción HTML directa'],
        ['Scraping Campofrío',      'Operativo', '105 productos, SPA con /p/ URLs'],
        ['Scraping El Pozo',        'Operativo', '493 productos, paginación 42 páginas'],
        ['Scraping Argal',          'Operativo', '148 productos, playwright-stealth'],
        ['Base de datos',           'Operativo', 'Supabase, 881 productos totales'],
        ['Detección de cambios',    'Operativo', 'Hash SHA-256 + comparación campo a campo'],
        ['Clasificación de calidad','Operativo', 'MEJORA / EMPEORAMIENTO / NEUTRO'],
        ['Alertas por email',       'Operativo', 'Gmail (local) + SendGrid (nube)'],
        ['Ejecución en la nube',    'Operativo', 'GitHub Actions, 2x/día + informe 9:15'],
        ['Dashboard visual',        'Operativo', 'HTML + Supabase SDK, responsive'],
        ['Análisis ingredientes',   'Operativo', 'Clasificación funcional automática'],
    ],
    col_widths_cm=[4.5, 2.5, 9]
)

heading2('3.2 Limitaciones actuales')
bullet('Datos nutricionales vacíos para El Pozo, Campofrío y Argal: estas webs muestran la tabla nutricional como imagen. Se requiere la API key de Claude Vision (Anthropic) para extraer estos datos. Coste estimado: 10–20 €/mes.')
bullet('Email remitente es una cuenta Gmail personal: se necesita la cuenta corporativa agente-etiquetado@tello.es para producción.')
bullet('Dashboard accesible solo localmente: se necesita publicar en Netlify o servidor interno para acceso compartido por el equipo.')
bullet('Repositorio GitHub en cuenta personal: pendiente migración a cuenta corporativa Tello.')

# ═══════════════════════════════════════════════════════════
# 4. PLAN HASTA 13 JULIO 2026
# ═══════════════════════════════════════════════════════════
doc.add_page_break()
heading1('4. PLAN DE TRABAJO HASTA EL 13 DE JULIO 2026')

heading2('Semana 4 — Validación y correcciones (10–14 Junio 2026)')
bullet('Verificar que los emails automáticos llegan correctamente cada día a las 9:15')
bullet('Revisar los datos en Supabase y corregir cualquier inconsistencia detectada')
bullet('Ajustar selectores CSS si alguna web de la competencia cambia su estructura')
bullet('Enviar informe de avance del proyecto a la responsable')
bullet('Solicitar a IT: API key de Claude Vision, cuenta email corporativa, validación GitHub')

heading2('Semana 5 — Mejoras del dashboard (17–21 Junio 2026)')
bullet('Implementar mapa de claims dinámico cargado desde la base de datos')
bullet('Mejorar pantalla nutricional con datos reales cuando estén disponibles')
bullet('Añadir exportación a Excel real (.xlsx) con datos completos')
bullet('Publicar dashboard en Netlify para acceso compartido por el equipo')
bullet('Crear URL permanente para compartir con el equipo de etiquetado')

heading2('Semana 6 — Integración Claude Vision (24–28 Junio 2026)')
bullet('Activar extracción de tablas nutricionales en imagen para El Pozo (condicionado a aprobación de API key)')
bullet('Activar extracción para Campofrío y Argal')
bullet('Verificar que los 881 productos tienen datos nutricionales completos')
bullet('Actualizar dashboard con los nuevos datos nutricionales')

heading2('Semana 7 — Configuración corporativa (1–5 Julio 2026)')
bullet('Migrar repositorio GitHub a cuenta de organización Tello')
bullet('Configurar cuenta agente-etiquetado@tello.es como remitente de los informes')
bullet('Actualizar destinatarios del informe diario con las personas definidas por el equipo')
bullet('Documentar proceso de mantenimiento para el equipo de IT')

heading2('Semana 8 — Cierre y entrega (8–13 Julio 2026)')
bullet('Demo del sistema completo al equipo de etiquetado')
bullet('Documentación técnica final para IT')
bullet('Manual de usuario del dashboard')
bullet('Revisión final con la responsable del proyecto')
bullet('Entrega formal del proyecto')

# ═══════════════════════════════════════════════════════════
# 5. NOTAS TÉCNICAS
# ═══════════════════════════════════════════════════════════
doc.add_page_break()
heading1('5. NOTAS TÉCNICAS')

heading2('5.1 Estructura del proyecto')
add_table(
    ['Archivo / Carpeta', 'Descripción'],
    [
        ['main.py',                  'Orquestador principal del ciclo de scraping'],
        ['config.py',               'Configuración centralizada (API keys, URLs, umbrales)'],
        ['calidad.py',              'Motor de clasificación de calidad de cambios'],
        ['alertas.py',              'Plantillas HTML de emails de alerta'],
        ['alertas_scheduler.py',    'Gestor de alertas programadas (informe diario + nuevos)'],
        ['scrapers/base_scraper.py','Clase base con Playwright, OCR y métodos compartidos'],
        ['scrapers/noel_scraper.py','Scraper específico de Noel'],
        ['scrapers/campofrio_scraper.py', 'Scraper específico de Campofrío'],
        ['scrapers/elpozo_scraper.py', 'Scraper específico de El Pozo'],
        ['scrapers/argal_scraper.py', 'Scraper específico de Argal (con stealth)'],
        ['database/schema.sql',     'Definición de tablas y vistas en PostgreSQL'],
        ['database/db.py',          'Capa de acceso a datos (UPSERT, historial, comparación)'],
        ['.github/workflows/',      'Automatización GitHub Actions (scraping + email en nube)'],
        ['dashboard_conectado.html','Interfaz visual conectada a Supabase en tiempo real'],
    ],
    col_widths_cm=[5.5, 10]
)

heading2('5.2 Infraestructura')
bullet('Repositorio: github.com/inessantistebanmoreno-dotcom/ia-competitiva-tello')
bullet('Base de datos: Supabase (plan gratuito) — ujassvznjpnfwtmpznsy.supabase.co')
bullet('Ejecución en la nube: GitHub Actions (gratuito con repositorio público)')
bullet('Email desde la nube: SendGrid (100 emails/día gratuitos)')
bullet('Dashboard local: archivo HTML abierto en Chrome (sin servidor necesario)')

heading2('5.3 Pendientes de resolución por IT')
bullet('API key de Claude Vision (Anthropic): necesaria para datos nutricionales en imagen. Coste ~10–20 €/mes')
bullet('Cuenta de correo corporativa agente-etiquetado@tello.es para envíos automáticos')
bullet('Validación del repositorio GitHub para alojamiento en cuenta corporativa Tello')
bullet('Servidor o URL interna para acceso compartido al dashboard')

# ── Guardar ──
output_path = r'C:\Users\temp03\Documents\Claude\Projects\Agente de análisis competencia\Memoria_Proyecto_Agente_Intel_Competitiva.docx'
doc.save(output_path)
print(f'Documento guardado en: {output_path}')
